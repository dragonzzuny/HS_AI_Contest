#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
화성ON 백엔드 — 다국어 외국인 생활민원 AI 안내사.

단일 FastAPI 앱: 프론트 서빙 + 인용기반 답변 + 처리보조(서류초안·체크리스트·딥링크).
LLM/벡터DB가 없어도 'MOCK 모드'로 즉시 작동(키워드 검색 + 템플릿 답변)하여 UX 검증 가능.
Ollama/Weaviate가 잡히면 자동으로 실동작(하이브리드 검색 + 그라운딩 생성)으로 승격.

실행:
  pip install -r backend/requirements.txt
  uvicorn backend.app:app --host 0.0.0.0 --port 5200
환경변수(선택):
  OLLAMA_HOST=http://localhost:11434   GEN_MODEL=qwen2.5:14b   EMBED_MODEL=bge-m3
  WEAVIATE_HOST=localhost
"""
import json
import os
import re
from pathlib import Path
from typing import Optional

import io
from urllib.parse import quote

import httpx
from fastapi import FastAPI
from fastapi.responses import JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

ROOT = Path(__file__).resolve().parent.parent
KB_JSON = ROOT / "kb" / "build" / "kb.json"
ORD_JSON = ROOT / "kb" / "build" / "ordinances_chunks.json"
ONTOLOGY_JSON = ROOT / "kb" / "build" / "ontology.json"
CACHE_JSON = ROOT / "kb" / "build" / "answer_cache.json"
WEB_DIR = ROOT / "web"

OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "").strip()      # 비면 MOCK 생성
GEN_MODEL = os.environ.get("GEN_MODEL", "qwen2.5:14b").strip()
EMBED_MODEL = os.environ.get("EMBED_MODEL", "bge-m3").strip()
WEAVIATE_HOST = os.environ.get("WEAVIATE_HOST", "").strip()  # 비면 KB 키워드 검색

LANGS = {
    "ko": "한국어", "en": "English", "zh": "中文",
    "vi": "Tiếng Việt", "th": "ภาษาไทย", "km": "ភាសាខ្មែរ",
}

# ── 처리보조: 도메인별 신청 딥링크 ─────────────────────────────
DEEPLINKS = {
    "체류": [
        {"label": "하이코리아 전자민원/방문예약", "url": "https://www.hikorea.go.kr"},
    ],
    "노무": [
        {"label": "고용노동부 노동포털(진정서 온라인 접수)", "url": "https://labor.moel.go.kr"},
    ],
    "건강보험": [
        {"label": "국민건강보험공단 민원여기요", "url": "https://www.nhis.or.kr"},
    ],
    "폐기물": [
        {"label": "대형폐기물 신고(빼기)", "url": "https://bbegi.com"},
        {"label": "폐가전 무상방문수거", "url": "https://www.15990903.or.kr"},
    ],
    "보육교육": [
        {"label": "아이사랑(어린이집 입소대기)", "url": "https://www.childcare.go.kr"},
    ],
    "운전면허": [
        {"label": "도로교통공단 안전운전 통합민원", "url": "https://www.safedriving.or.kr"},
    ],
    "행정": [
        {"label": "정부24", "url": "https://www.gov.kr"},
    ],
    "지원기관": [],
}

# ── 처리보조: 서류초안 입력 스키마(빈칸 → 한국어 공식 서류) ──────
FORM_SCHEMAS = {
    "노무": {
        "title": "임금체불 진정서",
        "desc": "고용노동부 노동포털 진정서 작성용 한국어 초안을 만듭니다.",
        "fields": [
            {"key": "name", "label": "이름", "ph": "Nguyen Van A"},
            {"key": "nationality", "label": "국적", "ph": "베트남"},
            {"key": "phone", "label": "연락처", "ph": "010-0000-0000"},
            {"key": "company", "label": "사업장명", "ph": "○○산업(주)"},
            {"key": "company_addr", "label": "사업장 주소", "ph": "화성시 향남읍 ..."},
            {"key": "start_date", "label": "입사일", "ph": "2025-03-01"},
            {"key": "end_date", "label": "퇴사일(재직 중이면 '재직')", "ph": "재직"},
            {"key": "amount", "label": "체불 금액(원)", "ph": "3000000"},
            {"key": "period", "label": "체불 기간", "ph": "2026-03 ~ 2026-05"},
        ],
    },
    "폐기물": {
        "title": "대형폐기물 배출 신고",
        "desc": "빼기(bbegi)/화성시 청소행정포털 신청용 내용을 정리합니다.",
        "fields": [
            {"key": "item", "label": "품목", "ph": "2인용 소파"},
            {"key": "qty", "label": "수량", "ph": "1"},
            {"key": "date", "label": "배출 희망일", "ph": "2026-06-25"},
            {"key": "location", "label": "배출 장소", "ph": "○○아파트 분리수거장"},
            {"key": "phone", "label": "연락처", "ph": "010-0000-0000"},
        ],
    },
}


def _f(fields, key, default="(미입력)"):
    v = (fields or {}).get(key, "")
    return v.strip() if isinstance(v, str) and v.strip() else default


def render_draft(domain, fields, today="년 월 일"):
    """입력값 → 한국어 공식 서류 초안(템플릿). LLM 없이도 동작."""
    if domain == "노무":
        amount = _f(fields, "amount", "")
        amount_s = f"{int(amount):,}" if amount.isdigit() else _f(fields, "amount")
        return (
            "진 정 서 (임금체불)\n\n"
            "1. 진정인\n"
            f"  - 성명: {_f(fields,'name')}\n"
            f"  - 국적: {_f(fields,'nationality')}\n"
            f"  - 연락처: {_f(fields,'phone')}\n\n"
            "2. 피진정인(사업주)\n"
            f"  - 사업장명: {_f(fields,'company')}\n"
            f"  - 사업장 소재지: {_f(fields,'company_addr')}\n\n"
            "3. 진정 내용\n"
            f"  본인은 {_f(fields,'company')}에서 "
            + (f"{_f(fields,'start_date')}부터 현재까지 재직 중이며, "
               if _f(fields, 'end_date') in ("재직", "재직중", "현재") else
               f"{_f(fields,'start_date')}부터 {_f(fields,'end_date')}까지 근무하였으며, ")
            + f"{_f(fields,'period')} 기간의 임금 합계 금 {amount_s}원을 "
            "정당한 사유 없이 지급받지 못하였기에, 「근로기준법」에 따라 "
            "체불 임금의 지급을 요청하고자 본 진정서를 제출합니다.\n\n"
            f"4. 요구사항: 체불 임금 {amount_s}원의 즉시 지급\n\n"
            f"{today}\n"
            f"진정인: {_f(fields,'name')} (서명)\n\n"
            "중부지방고용노동청 경기지청장 귀하\n\n"
            "※ 제출: 고용노동부 노동포털(labor.moel.go.kr) 온라인 진정 / 통역상담 ☎1350"
        )
    if domain == "폐기물":
        return (
            "[대형폐기물 배출 신고 내용]\n\n"
            f"  - 배출 품목: {_f(fields,'item')}\n"
            f"  - 수량: {_f(fields,'qty')}\n"
            f"  - 배출 희망일: {_f(fields,'date')}\n"
            f"  - 배출 장소: {_f(fields,'location')}\n"
            f"  - 신청인 연락처: {_f(fields,'phone')}\n\n"
            "※ 신청: 빼기(bbegi.com) 앱 또는 화성시 청소행정포털에 위 내용으로 신청 →\n"
            "  결제 후 발급된 예약번호를 폐기물에 부착하여 배출하세요.\n"
            "  문의: 화성시 자원순환과 ☎031-5189-6818"
        )
    return ""


HWPX_TEMPLATE = ROOT / "kb" / "hwpx_template"
FORMS_DIR = ROOT / "kb" / "forms"


def _xesc(t: str) -> str:
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# 실제 공식 서식(.hwpx) 등록부: { domain: {template, replace:{예시문구→"{필드}"}} }
def _load_form_registry():
    p = FORMS_DIR / "registry.json"
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


FORM_REGISTRY = _load_form_registry()


def fill_hwpx_template(path, replacements: dict) -> bytes:
    """실제 공식 .hwpx 서식의 <hp:t> 텍스트(예시값/빈칸)를 사용자 값으로 치환.
    한글 없이 서버에서 동작하며, 결과는 실제 한글에서 열림(검증 완료)."""
    import re
    import zipfile

    zin = zipfile.ZipFile(path)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zo:
        for it in zin.infolist():
            data = zin.read(it.filename)
            if it.filename == "mimetype":
                zi = zipfile.ZipInfo("mimetype")
                zi.compress_type = zipfile.ZIP_STORED
                zo.writestr(zi, data)
                continue
            if re.match(r"Contents/section\d+\.xml$", it.filename):
                s = data.decode("utf-8")

                def _f(m):
                    t = m.group(1)
                    for a, b in replacements.items():
                        t = t.replace(_xesc(a), _xesc(b))
                    return "<hp:t>" + t + "</hp:t>"

                s = re.sub(r"<hp:t>(.*?)</hp:t>", _f, s, flags=re.S)
                data = s.encode("utf-8")
            zo.writestr(it, data)
    return buf.getvalue()


def build_hwpx_bytes(text: str) -> bytes:
    """채워진 서류 텍스트 → 유효한 .hwpx(한글). 실제 한글이 만든 템플릿의 헤더·페이지설정을
    그대로 재사용하고 본문(section0)만 교체해 호환성을 확보한다."""
    import re
    import zipfile

    tdir = HWPX_TEMPLATE
    sec = (tdir / "Contents" / "section0.xml").read_text(encoding="utf-8")
    # 실제 한글이 만든 헤더(폰트·스타일)는 그대로 두고, 본문(section0)만 깨끗한 문단으로 교체.
    # 검증결과: 문단마다 <hp:linesegarray>가 있어야 한글이 정상적으로 연다(없으면 열기 실패).
    prefix = re.match(r"^<\?xml.*?<hs:sec[^>]*>", sec, re.S).group(0)
    lsa = ('<hp:linesegarray><hp:lineseg textpos="0" vertpos="0" vertsize="1700" '
           'textheight="1700" baseline="1445" spacing="1020" horzpos="0" '
           'horzsize="49324" flags="393216"/></hp:linesegarray>')

    paras = []
    for i, line in enumerate(text.split("\n")):
        paras.append(
            f'<hp:p id="{900001 + i}" paraPrIDRef="1" styleIDRef="0" '
            f'pageBreak="0" columnBreak="0" merged="0">'
            f'<hp:run charPrIDRef="0"><hp:t>{_xesc(line)}</hp:t></hp:run>{lsa}</hp:p>'
        )
    new_sec = prefix + "".join(paras) + "</hs:sec>"

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        zi = zipfile.ZipInfo("mimetype")           # mimetype은 반드시 첫 항목·무압축
        zi.compress_type = zipfile.ZIP_STORED
        z.writestr(zi, "application/hwp+zip")
        for p in sorted(tdir.rglob("*")):
            if not p.is_file():
                continue
            rel = p.relative_to(tdir).as_posix()
            if rel == "mimetype":
                continue
            if rel == "Contents/section0.xml":
                data = new_sec.encode("utf-8")
            elif rel == "Preview/PrvText.txt":
                data = text.encode("utf-8")
            else:
                data = p.read_bytes()
            z.writestr(rel, data)
    return buf.getvalue()


def build_docx_bytes(text: str, title: str) -> bytes:
    """채워진 서류 텍스트 → 한글 폰트 적용된 .docx 바이트(누구나 무료로 열고 편집 가능)."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    lines = text.split("\n")
    first_done = False
    for ln in lines:
        p = doc.add_paragraph()
        run = p.add_run(ln)
        if not first_done and ln.strip():       # 첫 실질 줄 = 제목
            run.bold = True
            run.font.size = Pt(15)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_done = True
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


class ChatRequest(BaseModel):
    message: str
    lang: str = "ko"


class DraftRequest(BaseModel):
    domain: str
    fields: dict = {}
    lang: str = "ko"
    today: str = "년    월    일"


def _tokens(s: str):
    s = (s or "").lower()
    words = [t for t in re.split(r"[^0-9a-z가-힣]+", s) if len(t) >= 2]
    # CJK(한자)는 단어경계가 없으므로 문자 bigram으로 보강
    cjk = re.findall(r"[一-鿿]", s)
    words += [cjk[i] + cjk[i + 1] for i in range(len(cjk) - 1)]
    return words


# 도메인별 다국어 트리거(부분일치) — 목 모드 교차언어 라우팅용.
# (실동작에서는 BGE-M3 교차언어 임베딩이 이 역할을 대체)
DOMAIN_TRIGGERS = {
    # 주의: 'register/đăng ký/ลงทะเบียน/ចុះឈ្មោះ' 같은 일반 동사는 다른 도메인(건강보험 가입 등)과
    # 충돌하므로 제외하고, 외국인·비자·체류 등 도메인 특이어만 트리거로 사용.
    "체류": ["외국인등록", "체류", "비자", "등록증", "foreigner", "alien", "visa",
            "residence", "外国人", "登录", "居留", "签证", "người nước ngoài",
            "thị thực", "cư trú", "ต่างชาติ", "วีซ่า",
            "ជនបរទេស", "ទិដ្ឋាការ"],
    "폐기물": ["쓰레기", "분리수거", "종량제", "대형폐기물", "음식물", "재활용",
             "trash", "garbage", "waste", "rubbish", "recycle", "垃圾", "扔垃圾", "分类",
             "rác", "vứt rác", "ขยะ", "ทิ้งขยะ", "សំរាម", "បោះសំរាម"],
    "노무": ["임금", "월급", "급여", "체불", "근로", "노동", "산재", "퇴직금",
            "wage", "salary", "unpaid", "labor", "labour", "work", "工资", "拖欠", "劳动",
            "lương", "tiền lương", "nợ lương", "ค่าจ้าง", "เงินเดือน", "ប្រាក់ឈ្នួល"],
    "건강보험": ["건강보험", "보험료", "병원", "의료", "health insurance", "premium", "hospital",
              "medical", "健康保险", "医院", "保险", "bảo hiểm", "bệnh viện",
              "ประกันสุขภาพ", "โรงพยาบาล", "ធានារ៉ាប់រង", "មន្ទីរពេទ្យ"],
    "보육교육": ["어린이집", "유치원", "학교", "입학", "보육", "교육", "자녀",
              "daycare", "kindergarten", "school", "enroll", "child", "幼儿园", "学校", "入学",
              "nhà trẻ", "trường học", "โรงเรียน", "อนุบาล", "សាលា", "កុមារ"],
    "운전면허": ["운전면허", "면허", "운전", "driver", "license", "licence", "driving",
              "驾照", "驾驶", "bằng lái", "lái xe", "ใบขับขี่", "ប័ណ្ណបើកបរ"],
    "행정": ["전입신고", "주소", "이사", "주민센터", "move", "address", "搬家", "地址",
            "chuyển nhà", "địa chỉ", "ย้ายบ้าน", "ផ្លាស់ប្ដូរ"],
    "지원기관": ["상담", "통역", "가족센터", "지원센터", "외국인주민", "counsel", "counseling",
              "support center", "interpret", "咨询", "翻译", "支援", "tư vấn", "hỗ trợ", "phiên dịch",
              "ปรึกษา", "ช่วยเหลือ", "ล่าม", "ប្រឹក្សា", "ជំនួយ", "បកប្រែ"],
}


def valid_ord_content(s: str) -> bool:
    """수집 실패/오류 레코드(쓰레기 청크) 판별 — 검색·인용에서 제외(환각 방지)."""
    s = (s or "").strip()
    if not s:
        return False
    if "일치하는 자치법규가 없" in s:            # 법제처 API 오류 응답
        return False
    if s.startswith("{") and '"Law"' in s[:40]:  # JSON 오류 페이로드
        return False
    return True


class Corpus:
    """KB(L1) + 자치법규(L2) 로드 및 키워드 검색(폴백)."""
    def __init__(self):
        self.items = []
        self.dropped_ord = 0
        if KB_JSON.exists():
            for it in json.loads(KB_JSON.read_text(encoding="utf-8")):
                self.items.append({
                    "id": it["id"], "layer": "L1",
                    "title": it.get("title", ""), "domain": it.get("domain", ""),
                    "answer": it.get("answer", ""), "target": it.get("target", ""),
                    "documents": it.get("documents", ""), "dept": it.get("dept", ""),
                    "phone": it.get("phone", ""), "location": it.get("location", ""),
                    "source": it.get("source", ""),
                    "confidence": (it.get("confidence", "") or "").split()[0],
                    "questions": it.get("questions", []),
                    "_blob": " ".join([it.get("title", ""), it.get("target", ""),
                                       it.get("answer", ""), it.get("documents", ""),
                                       " ".join(it.get("questions", []))]),
                })
        if ORD_JSON.exists():
            for c in json.loads(ORD_JSON.read_text(encoding="utf-8")):
                if not valid_ord_content(c.get("content", "")):
                    self.dropped_ord += 1
                    continue
                self.items.append({
                    "id": c["id"], "layer": "L2",
                    "title": f"{c.get('law_name','')} {c.get('article','')}",
                    "domain": "법령", "answer": c.get("content", ""), "target": "",
                    "documents": "", "dept": c.get("kind", ""), "phone": "",
                    "location": "", "source": c.get("source", ""),
                    "confidence": "확인됨", "questions": [],
                    "_blob": f"{c.get('law_name','')} {c.get('content','')}",
                })

    def search(self, query: str, k: int = 4):
        ql = (query or "").lower()
        q = set(_tokens(query))
        # 법령 의도가 명시된 질의만 L2(조례)를 상위 노출
        law_intent = any(w in ql for w in
                         ["조례", "규칙", "법령", "조문", "자치법규", "law", "ordinance", "article"])
        # 다국어 트리거로 매칭된 도메인(교차언어 라우팅)
        boosted = {dom for dom, kws in DOMAIN_TRIGGERS.items()
                   if any(kw.lower() in ql for kw in kws)}
        scored = []
        for it in self.items:
            blob = set(_tokens(it["_blob"]))
            overlap = len(q & blob)
            if it["layer"] == "L1":
                score = overlap * 1.5
                for qx in it.get("questions", []):
                    if q & set(_tokens(qx)):
                        score += 3
                if it["domain"] in boosted:
                    score += 10          # 교차언어: 같은 도메인 L1 강하게 부스팅
            else:  # L2 자치법규
                score = overlap * (0.8 if law_intent else 0.25)
                if it["domain"] in boosted:
                    score += 1
            if score > 0:
                it["_score"] = round(score, 2)
                scored.append((score, it))
        scored.sort(key=lambda x: x[0], reverse=True)
        ordered = [it for _, it in scored]
        if not law_intent:  # 생활민원 질의는 L1을 항상 우선
            ordered = [it for it in ordered if it["layer"] == "L1"] + \
                      [it for it in ordered if it["layer"] == "L2"]
        return ordered[:k]


CORPUS = Corpus()
app = FastAPI(title="화성ON")

# ── 환각(헛소리) 방지 가드레일 ────────────────────────────────
# 근거 신뢰도가 기준 미만이면 '단정'하지 않고 담당기관으로 안전 안내.
# L1 트리거 매칭(+10)·질문 매칭(+3)은 통과, 일반 키워드 잡음(<7)은 거부.
GATE_L1 = 7.0     # 생활민원 근거 최소 점수
GATE_L2 = 2.0     # 명시적 법령질의 시 조문 근거 최소 점수
CITE_RE = re.compile(r"\[([a-z]+-[a-z]+-\d+|ord-[\w-]+)\]")

# 근거를 찾지 못했을 때의 안전 안내(언어별). 절대 사실을 지어내지 않는다.
SAFE_REFUSAL = {
    "ko": ("죄송합니다. 화성시 생활민원에 한해 공식 자료에 근거가 있을 때만 답변드립니다. "
           "이 질문은 근거를 찾지 못해 정확히 안내드리기 어렵습니다.\n\n"
           "📞 화성시 외국인 종합상담 1688-0911 · 다누리콜센터 1577-1366(다국어)으로 문의해 주세요."),
    "en": ("Sorry — I only answer Hwaseong living/civil-affairs questions that are backed by official "
           "sources. I couldn't find a grounded source for this, so I won't guess.\n\n"
           "📞 Please contact Hwaseong foreign-resident help 1688-0911 or the Danuri call center "
           "1577-1366 (multilingual)."),
    "zh": ("抱歉,我仅就有华城市官方资料依据的生活民愿问题作答。此问题未找到依据,无法准确解答。\n\n"
           "📞 请拨打 华城市外国人综合咨询 1688-0911 或 Danuri 呼叫中心 1577-1366(多语种)。"),
    "vi": ("Xin lỗi — tôi chỉ trả lời các câu hỏi dân sinh ở Hwaseong khi có căn cứ từ nguồn chính "
           "thức. Tôi không tìm thấy căn cứ cho câu hỏi này nên sẽ không suy đoán.\n\n"
           "📞 Vui lòng liên hệ Tư vấn người nước ngoài Hwaseong 1688-0911 hoặc tổng đài Danuri "
           "1577-1366 (đa ngôn ngữ)."),
    "th": ("ขออภัย ฉันตอบเฉพาะคำถามด้านบริการประชาชนของฮวาซองที่มีข้อมูลทางการรองรับเท่านั้น "
           "คำถามนี้ไม่พบแหล่งอ้างอิงจึงจะไม่เดา\n\n"
           "📞 โปรดติดต่อศูนย์ให้คำปรึกษาชาวต่างชาติฮวาซอง 1688-0911 หรือสายด่วน Danuri 1577-1366 (หลายภาษา)"),
    "km": ("សុំទោស — ខ្ញុំឆ្លើយតែសំណួរសេវាប្រជាពលរដ្ឋ Hwaseong ដែលមានមូលដ្ឋានពីប្រភពផ្លូវការប៉ុណ្ណោះ។ "
           "សំណួរនេះរកមូលដ្ឋានមិនឃើញ ដូច្នេះខ្ញុំនឹងមិនស្មាន។\n\n"
           "📞 សូមទាក់ទង ការប្រឹក្សាជនបរទេស Hwaseong 1688-0911 ឬ មជ្ឈមណ្ឌល Danuri 1577-1366 (ពហុភាសា)។"),
}


def _law_intent(q: str) -> bool:
    ql = (q or "").lower()
    return any(w in ql for w in
               ["조례", "규칙", "법령", "조문", "자치법규", "law", "ordinance", "article"])


def is_grounded(query: str, hits) -> bool:
    """근거 신뢰도 게이트: 통과해야만 답변 생성. 미통과 시 안전 거부."""
    if not hits:
        return False
    l1 = [h for h in hits if h["layer"] == "L1"]
    if l1 and l1[0].get("_score", 0) >= GATE_L1:
        return True
    if _law_intent(query) and hits[0]["layer"] == "L2" and hits[0].get("_score", 0) >= GATE_L2:
        return True
    return False


def validate_citations(answer: str, hits) -> str:
    """답변 본문에 실제 근거 목록에 없는 [id]가 있으면 제거(유령 인용 방지)."""
    valid = {h["id"] for h in hits}
    return CITE_RE.sub(lambda m: m.group(0) if m.group(1) in valid else "", answer)


# ── 온톨로지: 도메인↔법적근거(자치법규)↔자격↔선행절차 ──────────────
ONTOLOGY = {}
if ONTOLOGY_JSON.exists():
    ONTOLOGY = json.loads(ONTOLOGY_JSON.read_text(encoding="utf-8"))

# ── 답변 캐시: 자주 묻는 질문은 미리 만든 답을 즉시 제공(언어별) ──────
# 키 = L1 근거 항목 ID(온톨로지 토픽). 흔한 질문은 캐시 즉답, 그 외는 LLM 실시간.
# 미리 만들어 커밋하면 GPU 없는 배포에서도 6개 언어 즉시 응답.
ANSWER_CACHE = {}
_CACHE_MTIME = [0.0]


def _reload_cache():
    """캐시 파일이 바뀌면(프리필 등) 메모리에 다시 로드 — 재시작 불필요."""
    try:
        m = CACHE_JSON.stat().st_mtime
    except OSError:
        return
    if m != _CACHE_MTIME[0]:
        try:
            data = json.loads(CACHE_JSON.read_text(encoding="utf-8"))
            ANSWER_CACHE.clear()
            ANSWER_CACHE.update(data)
            _CACHE_MTIME[0] = m
        except Exception:
            pass


_reload_cache()


def _cache_get(item_id, lang):
    _reload_cache()
    return (ANSWER_CACHE.get(item_id) or {}).get(lang)


def _cache_put(item_id, lang, answer):
    _reload_cache()                          # 다른 프로세스가 채운 것 먼저 반영(덮어쓰기 방지)
    ANSWER_CACHE.setdefault(item_id, {})[lang] = answer
    try:
        CACHE_JSON.write_text(json.dumps(ANSWER_CACHE, ensure_ascii=False, indent=1), encoding="utf-8")
        _CACHE_MTIME[0] = CACHE_JSON.stat().st_mtime
    except Exception:
        pass


# ── 공식 서식 카탈로그(고빈도·안정 공문서 수집) ────────────────────
CATALOG = {}
_cat = FORMS_DIR / "catalog.json"
if _cat.exists():
    try:
        CATALOG = json.loads(_cat.read_text(encoding="utf-8"))
    except Exception:
        CATALOG = {}


def catalog_form_actions(hits):
    """매칭 도메인의 공식 서식을 안내(수집/온라인 여부 표시)."""
    dom = top_domain(hits)
    if not dom:
        return []
    out = []
    for f in CATALOG.get("forms", []):
        if f.get("domain") != dom:
            continue
        st = f.get("status")
        tag = ("· 수집됨" if st == "collected" else "· 온라인 신청" if st == "online_only" else "· 공식 출처")
        ml = " · 다국어" if f.get("multilingual") else ""
        out.append({"type": "link",
                    "label": f"📄 공식 서식: {f['name']} ({f.get('issuer','')}) {tag}{ml}",
                    "url": f.get("source", "")})
    return out[:2]


def _loc(v, lang):
    """다국어 dict({ko,en,...})에서 언어 선택, 없으면 ko/en 폴백."""
    if isinstance(v, dict):
        return v.get(lang) or v.get("en") or v.get("ko") or ""
    return v or ""


def top_domain(hits):
    for h in hits:
        if h["layer"] == "L1" and h.get("domain"):
            return h["domain"]
    return None


def ontology_legal_citations(query, hits):
    """매칭 도메인의 법적 근거 중 '질문과 관련된 실효 조문'을 우선 인용(동적 선택)."""
    dom = top_domain(hits)
    node = (ONTOLOGY.get("domains") or {}).get(dom)
    if not node:
        return []
    have = {h["id"] for h in hits}
    qtok = set(_tokens(query))
    ranked = sorted(
        node.get("legal_basis", []),
        key=lambda lb: len(qtok & set(_tokens(lb.get("content", "")))),
        reverse=True,
    )
    out = []
    for lb in ranked:
        if lb["id"] in have:
            continue
        out.append({
            "id": lb["id"], "title": f"{lb['law_name']} {lb.get('article','')}".strip(),
            "layer": "L2", "snippet": lb.get("content", ""),
            "dept": "화성시 자치법규", "phone": "",
            "source": lb.get("source", ""), "confidence": "확인됨",
            "kind": "법적근거",
        })
        if len(out) >= 2:
            break
    return out


def ontology_policy_action(hits, lang):
    """매칭 도메인의 지원정책(혜택)을 안내 카드로 노출 — '받을 수 있는 지원'."""
    dom = top_domain(hits)
    node = (ONTOLOGY.get("domains") or {}).get(dom)
    if not node:
        return None
    pols = node.get("support_policies", [])
    if not pols:
        return None
    items = [{"name": p.get("name", ""), "desc": _loc(p.get("desc"), lang), "url": p.get("url", "")}
             for p in pols[:4]]
    title = "받을 수 있는 지원" if lang == "ko" else "Support you may qualify for"
    return {"type": "policy", "title": title, "items": items}


def ontology_info_action(hits, lang):
    """자격·선행절차·근거법령을 요약한 안내 카드(온톨로지 추론 노출)."""
    dom = top_domain(hits)
    node = (ONTOLOGY.get("domains") or {}).get(dom)
    if not node:
        return None
    items = []
    elig = _loc(node.get("eligibility"), lang)
    if elig:
        items.append(("대상" if lang == "ko" else "Eligibility") + f": {elig}")
    # 이 도메인을 가리키는 선행절차 관계
    for rel in ONTOLOGY.get("relations", []):
        if rel.get("to") == dom and rel.get("type") == "prerequisite_for":
            note = _loc(rel.get("note"), lang)
            if note:
                items.append(("선행" if lang == "ko" else "First") + f": {note}")
    nat = node.get("national_law")
    if nat:
        items.append(("근거법" if lang == "ko" else "Law") + f": {_loc(nat, lang)}")
    if not items:
        return None
    title = "자격·절차 안내" if lang == "ko" else "Eligibility & steps"
    return {"type": "info", "title": title, "items": items}


def build_actions(hits):
    """검색된 항목에서 처리보조(딥링크·체크리스트·서류초안) 구성."""
    actions = []
    seen = set()
    for h in hits:
        dom = h["domain"]
        for dl in DEEPLINKS.get(dom, []):
            if dl["url"] not in seen:
                actions.append({"type": "link", **dl})
                seen.add(dl["url"])
        if h.get("documents"):
            actions.append({"type": "checklist", "title": f"{h['title']} 필요서류",
                            "items": [x.strip() for x in re.split(r"[,·]", h["documents"]) if x.strip()]})
        if dom in FORM_SCHEMAS and not any(a["type"] == "form" for a in actions):
            sc = FORM_SCHEMAS[dom]
            actions.append({"type": "form", "domain": dom,
                            "title": sc["title"], "desc": sc["desc"], "fields": sc["fields"]})
    return actions[:6]


GROUNDED_SYS = (
    "당신은 화성시 거주 외국인을 돕는 '화성ON' 민원 안내사다. "
    "반드시 아래 제공된 [근거]만 사용해 답하라. 근거에 없는 내용은 지어내지 말고 "
    "'정확한 확인을 위해 담당기관에 문의하세요'라고 안내하라. "
    "답변은 사용자 언어({lang})로 쓰되, 핵심 절차는 단계별로, 마지막에 담당부서·전화를 제시하라. "
    "사용한 근거의 ID를 문장 끝에 [id] 형태로 표기하라."
)


def make_context(hits):
    blocks = []
    for h in hits:
        blocks.append(
            f"[{h['id']}] {h['title']}\n대상:{h.get('target','')}\n"
            f"내용:{h['answer']}\n담당:{h.get('dept','')} 전화:{h.get('phone','')}\n"
            f"출처:{h.get('source','')} 신뢰도:{h.get('confidence','')}"
        )
    return "\n\n".join(blocks)


def mock_answer(query, lang, hits):
    if not hits:
        return ("죄송합니다. 해당 민원 정보를 찾지 못했습니다. "
                "화성시 콜센터 ☎1688-0911 또는 다누리콜센터 ☎1577-1366으로 문의하세요. "
                "(MOCK 모드 — Ollama 연결 시 다국어 실답변)")
    top = hits[0]
    note = "" if lang == "ko" else f"\n\n[{LANGS.get(lang, lang)} 답변은 Ollama 연결 시 제공됩니다 — 현재 MOCK]"
    dept, phone = top.get("dept", ""), top.get("phone", "")
    meta = f"📌 담당: {dept}" + (f" {phone}" if phone else "")   # phone에 이미 ☎ 포함
    return f"{top['answer']}\n\n{meta} [{top['id']}]{note}"


async def ollama_generate(query, lang, hits) -> Optional[str]:
    if not OLLAMA_HOST:
        return None
    prompt = (f"[근거]\n{make_context(hits)}\n\n[질문]\n{query}\n\n"
              f"위 근거만으로 {LANGS.get(lang,'한국어')}로 답하라.")
    try:
        async with httpx.AsyncClient(timeout=120) as cli:
            r = await cli.post(f"{OLLAMA_HOST}/api/chat", json={
                "model": GEN_MODEL, "stream": False,
                "messages": [
                    {"role": "system", "content": GROUNDED_SYS.format(lang=LANGS.get(lang, "한국어"))},
                    {"role": "user", "content": prompt},
                ],
            })
            r.raise_for_status()
            return r.json().get("message", {}).get("content", "").strip()
    except Exception as e:
        return f"[Ollama 오류: {e}] " + mock_answer(query, lang, hits)


@app.post("/api/draft")
async def draft(req: DraftRequest):
    base = render_draft(req.domain, req.fields, req.today)
    if not base:
        return JSONResponse({"draft": "", "error": "지원하지 않는 서식입니다."}, status_code=400)
    # Ollama가 있으면 자연스럽게 다듬고(한국어 서식 유지) + 사용자 언어 요약 병기
    if OLLAMA_HOST and req.lang != "ko":
        try:
            async with httpx.AsyncClient(timeout=120) as cli:
                r = await cli.post(f"{OLLAMA_HOST}/api/chat", json={
                    "model": GEN_MODEL, "stream": False,
                    "messages": [
                        {"role": "system", "content":
                            "다음 한국어 공식 서류 초안을 그대로 유지하되, 맨 아래에 "
                            f"'--- {LANGS.get(req.lang)} 요약 ---' 구분선과 함께 {LANGS.get(req.lang)}로 "
                            "핵심 내용을 3줄 요약해 덧붙여라. 한국어 본문은 절대 바꾸지 마라."},
                        {"role": "user", "content": base},
                    ],
                })
                r.raise_for_status()
                polished = r.json().get("message", {}).get("content", "").strip()
                if polished:
                    return JSONResponse({"draft": polished, "mode": "ollama"})
        except Exception:
            pass
    return JSONResponse({"draft": base, "mode": "mock"})


def _fmt_fields(domain, fields):
    """서식 채움용 필드값(금액 천단위 등 서식 적용)."""
    f = dict(fields or {})
    amt = str(f.get("amount", "")).replace(",", "")
    if amt.isdigit():
        f["amount_won"] = f"{int(amt):,}원"
        f["amount"] = f"{int(amt):,}"
    return {k: (v if isinstance(v, str) else str(v)) for k, v in f.items()}


@app.post("/api/draft/hwpx")
async def draft_hwpx(req: DraftRequest):
    title = FORM_SCHEMAS.get(req.domain, {}).get("title", "서류초안")
    fname = f"{title.replace(' ', '_')}.hwpx"

    # 1) 실제 공식 .hwpx 서식이 등록돼 있으면 그 서식의 빈칸을 채움(우선).
    reg = FORM_REGISTRY.get(req.domain)
    if reg and (FORMS_DIR / reg["template"]).exists():
        ff = _fmt_fields(req.domain, req.fields)
        safe = {k: ff.get(k, "") for k in set(re.findall(r"{(\w+)}", json.dumps(reg.get("replace", {}))))}
        reps = {}
        for example, tpl in reg.get("replace", {}).items():
            try:
                reps[example] = tpl.format_map({**{k: "" for k in safe}, **ff})
            except Exception:
                reps[example] = tpl
        data = fill_hwpx_template(FORMS_DIR / reg["template"], reps)
        return Response(content=data, media_type="application/hwp+zip",
                        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(fname)}"})

    # 2) 등록된 공식 서식이 없으면 검증된 생성본으로 폴백(한글에서 열림 확인됨).
    base = render_draft(req.domain, req.fields, req.today)
    if not base:
        return JSONResponse({"error": "지원하지 않는 서식입니다."}, status_code=400)
    if not HWPX_TEMPLATE.exists():
        return JSONResponse({"error": "HWPX 템플릿이 없습니다."}, status_code=500)
    data = build_hwpx_bytes(base)
    return Response(content=data, media_type="application/hwp+zip",
                    headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(fname)}"})


@app.post("/api/draft/docx")
async def draft_docx(req: DraftRequest):
    base = render_draft(req.domain, req.fields, req.today)
    if not base:
        return JSONResponse({"error": "지원하지 않는 서식입니다."}, status_code=400)
    title = FORM_SCHEMAS.get(req.domain, {}).get("title", "서류초안")
    data = build_docx_bytes(base, title)
    fname = f"{title.replace(' ', '_')}.docx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(fname)}"},
    )


@app.get("/api/health")
def health():
    return {
        "ok": True,
        "corpus": len(CORPUS.items),
        "l1": sum(1 for i in CORPUS.items if i["layer"] == "L1"),
        "l2": sum(1 for i in CORPUS.items if i["layer"] == "L2"),
        "l2_dropped_garbage": CORPUS.dropped_ord,
        "ontology_domains": len((ONTOLOGY.get("domains") or {})),
        "ontology_legal_links": sum(len(d.get("legal_basis", []))
                                    for d in (ONTOLOGY.get("domains") or {}).values()),
        "mode": "ollama" if OLLAMA_HOST else "mock",
        "gen_model": GEN_MODEL if OLLAMA_HOST else None,
        "langs": LANGS,
    }


@app.post("/api/chat")
async def chat(req: ChatRequest):
    hits = CORPUS.search(req.message, k=4)

    # 환각 방지: 근거 신뢰도 미달이면 답변을 생성하지 않고 안전 안내.
    if not is_grounded(req.message, hits):
        lang = req.lang if req.lang in SAFE_REFUSAL else "ko"
        return JSONResponse({
            "answer": SAFE_REFUSAL[lang],
            "citations": [], "actions": [],
            "grounded": False, "refused": True,
            "mode": "ollama" if OLLAMA_HOST else "mock",
        })

    # 하이브리드: ① 캐시(미리 만든 답) 즉답 → ② LLM 실시간(성공 시 캐시 저장) → ③ MOCK 폴백
    top_l1 = next((h for h in hits if h["layer"] == "L1"), hits[0])
    cache_key = top_l1["id"]
    served = "cache"
    answer = _cache_get(cache_key, req.lang)
    if answer is None:
        gen = await ollama_generate(req.message, req.lang, hits)
        if gen is None:                     # OLLAMA 미연결 → MOCK(한국어 즉답)
            answer = mock_answer(req.message, req.lang, hits)
            served = "mock"
        elif "[Ollama 오류" in gen:          # LLM 오류 → 폴백(캐시 안 함)
            answer = gen
            served = "llm_error"
        else:                                # LLM 실답변 → 캐시에 저장(다음엔 즉답)
            answer = gen
            served = "llm"
            _cache_put(cache_key, req.lang, gen)
    answer = validate_citations(answer, hits)   # 유령 인용 제거

    citations = [{
        "id": h["id"], "title": h["title"], "layer": h["layer"],
        "snippet": (h["answer"][:200] + ("…" if len(h["answer"]) > 200 else "")),
        "dept": h.get("dept", ""), "phone": h.get("phone", ""),
        "source": h.get("source", ""), "confidence": h.get("confidence", ""),
        "score": h.get("_score", 0),
    } for h in hits]
    # 온톨로지 보강: 질문 관련 법적 근거(자치법규) 인용 + 자격·선행절차 + 지원정책
    citations += ontology_legal_citations(req.message, hits)
    actions = build_actions(hits)
    info = ontology_info_action(hits, req.lang)
    policy = ontology_policy_action(hits, req.lang)
    head = [a for a in (info, policy) if a]
    actions = head + catalog_form_actions(hits) + actions
    return JSONResponse({
        "answer": answer,
        "citations": citations,
        "actions": actions[:9],
        "grounded": True, "refused": False,
        "served": served,
        "mode": "ollama" if OLLAMA_HOST else "mock",
    })


# 프론트 정적 서빙 (맨 마지막에 마운트)
if WEB_DIR.exists():
    app.mount("/", StaticFiles(directory=str(WEB_DIR), html=True), name="web")
